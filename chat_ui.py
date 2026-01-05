"""chat_ui.py

Interfaz web (Gradio) para un asistente RAG sobre ChromaDB + LlamaIndex + Gemini.

Incluye:
- Entrada por teclado: Enter envía, Shift+Enter añade salto de línea.
- Gestión de sesiones persistentes (crear, renombrar, borrar, limpiar) con dropdown robusto.
- Exportación de la conversación actual a .docx con conversión Markdown -> Word básica.
- Ocultación del botón de compartir del Chatbot (si la versión de Gradio lo soporta).

Nota: Se mantienen intactas la lógica y las dependencias del pipeline RAG.
"""

import os
import json
import uuid
import time
import traceback
import ast
import re
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional, Iterable
from concurrent.futures import ThreadPoolExecutor

from dotenv import load_dotenv
import gradio as gr
import chromadb

from llama_index.core import Settings, VectorStoreIndex
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.llms.google_genai import GoogleGenAI
from llama_index.embeddings.google_genai import GoogleGenAIEmbedding


# -----------------------------------------------------------------------------
# Carga de variables de entorno
# -----------------------------------------------------------------------------
load_dotenv()


# -----------------------------------------------------------------------------
# Configuración (debe alinearse con ingest.py)
# -----------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent

CHROMA_PATH = os.getenv("CHROMA_PATH", str(BASE_DIR / "cerebro_db"))
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "tfg_master")

TOP_K_OPTIONS = [5, 10, 20]
DEFAULT_TOP_K = 10

GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3-pro-preview")
GEMINI_EMBED_MODEL = os.getenv("GEMINI_EMBED_MODEL", "text-embedding-004")

SESSIONS_DIR = BASE_DIR / "ui_sessions"
SESSIONS_FILE = SESSIONS_DIR / "sessions.json"
EXPORTS_DIR = SESSIONS_DIR / "exports"

# Límite opcional para reducir carga en UI (0 = sin límite)
MAX_SESSIONS = 0  # p. ej. 200 si se quiere mantener un histórico acotado

INSTRUCCIONES = """
Eres el asistente del TFG de Saúl sobre GNSS, DGNSS y HAPS.

REGLAS IMPORTANTES:
- Usa SIEMPRE la información de los documentos recuperados como fuente principal.
- Si la respuesta requiere varios documentos, intégralos y deja claro de dónde viene cada idea.
- Si no encuentras información suficiente en los documentos, dilo claramente y, si procede, explica qué información faltaría.
- No inventes fórmulas, valores numéricos ni conclusiones que no aparezcan (directa o razonablemente) en el contexto.
- Responde de forma estructurada (títulos, listas, pasos) cuando la pregunta sea compleja.
""".strip()


# -----------------------------------------------------------------------------
# Utilidades: tiempo, normalización y migración de históricos
# -----------------------------------------------------------------------------

def _now_ts() -> float:
    return float(time.time())


def _to_float_ts(v: Any) -> float:
    try:
        return float(v)
    except Exception:
        return 0.0


def _sanitize_title(title: str) -> str:
    """Normaliza títulos y elimina sufijos del tipo "(123456789)" cuando son ruido."""
    t = (title or "").strip()
    if not t:
        return "Conversación"
    if t.endswith(")") and "(" in t:
        base, tail = t.rsplit("(", 1)
        tail = tail[:-1].strip()  # sin ')'
        if tail.isdigit() and len(tail) >= 6:
            t = base.strip()
    return t[:80] if len(t) > 80 else t


def _content_parts_to_text(parts: Iterable[Any]) -> str:
    """Une una lista de 'parts' (formatos heterogéneos) en un único texto."""
    chunks: List[str] = []
    for p in parts:
        if p is None:
            continue
        if isinstance(p, str):
            chunks.append(p)
            continue
        if isinstance(p, dict):
            if "text" in p and isinstance(p.get("text"), (str, int, float)):
                chunks.append(str(p.get("text")))
                continue
            if p.get("type") == "text" and "content" in p:
                chunks.append(str(p.get("content") or ""))
                continue
            if "content" in p and isinstance(p.get("content"), str):
                chunks.append(p["content"])
                continue
            continue

        txt = getattr(p, "text", None)
        if isinstance(txt, str):
            chunks.append(txt)
            continue
        txt = getattr(p, "content", None)
        if isinstance(txt, str):
            chunks.append(txt)
            continue

        chunks.append(str(p))

    return "".join(chunks).strip()


def _content_to_text(content: Any) -> str:
    """Convierte diferentes representaciones de contenido a un string estable."""
    if content is None:
        return ""
    if isinstance(content, str):
        s = content
        s_strip = s.strip()
        # Corrección de históricos antiguos: strings que representaban 'parts' serializados.
        if (s_strip.startswith("[{") or s_strip.startswith("{")) and (
            ("'text'" in s_strip or '"text"' in s_strip)
            and ("'type'" in s_strip or '"type"' in s_strip)
        ):
            try:
                parsed = ast.literal_eval(s_strip)
                return _content_to_text(parsed)
            except Exception:
                pass
        return s
    if isinstance(content, (int, float, bool)):
        return str(content)
    if isinstance(content, (list, tuple)):
        return _content_parts_to_text(content)
    if isinstance(content, dict):
        if isinstance(content.get("text"), (str, int, float)):
            return str(content.get("text"))
        if isinstance(content.get("content"), (str, int, float)):
            return str(content.get("content"))
        return ""
    maybe = getattr(content, "text", None)
    if isinstance(maybe, str):
        return maybe
    maybe = getattr(content, "content", None)
    if isinstance(maybe, str):
        return maybe
    if isinstance(maybe, (list, tuple)):
        return _content_parts_to_text(maybe)
    return str(content)


def _extract_role_content(msg: Any) -> Tuple[str, Any]:
    """Extrae (role, content) de dicts o estructuras tipo ChatMessage."""
    if msg is None:
        return "user", ""
    if isinstance(msg, dict):
        return str(msg.get("role") or "user").lower().strip(), msg.get("content", "")
    role = getattr(msg, "role", None)
    content = getattr(msg, "content", None)
    if role is not None:
        return str(role).lower().strip(), content
    return "user", str(msg)


def _normalize_history_to_messages(hist: Any) -> List[Dict[str, str]]:
    """Normaliza el historial a una lista de mensajes con forma:

    [{"role":"user","content":"..."}, {"role":"assistant","content":"..."}, ...]

    Soporta:
    - Pares [[user, assistant], ...]
    - Lista de dicts con role/content
    - Lista de ChatMessage
    - Mezclas razonables
    """
    if not isinstance(hist, list) or not hist:
        return []

    # Caso: pares (formato clásico de Gradio)
    if isinstance(hist[0], (list, tuple)) and len(hist[0]) == 2:
        out: List[Dict[str, str]] = []
        for pair in hist:
            try:
                u, a = pair
                out.append({"role": "user", "content": _content_to_text(u)})
                out.append({"role": "assistant", "content": _content_to_text(a)})
            except Exception:
                continue
        return out

    out: List[Dict[str, str]] = []
    pending_user: Optional[str] = None

    for raw in hist:
        role, content_raw = _extract_role_content(raw)
        role = (role or "user").lower().strip()
        if role not in ("user", "assistant", "system"):
            role = "user"

        content = _content_to_text(content_raw)

        if role == "user":
            if pending_user is not None:
                out.append({"role": "user", "content": pending_user})
                out.append({"role": "assistant", "content": ""})
            pending_user = content
        elif role == "assistant":
            if pending_user is None:
                out.append({"role": "user", "content": ""})
                out.append({"role": "assistant", "content": content})
            else:
                out.append({"role": "user", "content": pending_user})
                out.append({"role": "assistant", "content": content})
                pending_user = None
        else:
            out.append({"role": "system", "content": content})

    if pending_user is not None:
        out.append({"role": "user", "content": pending_user})
        out.append({"role": "assistant", "content": ""})

    return out


def _migrate_sessions_store(store: Dict[str, Any]) -> Dict[str, Any]:
    """Asegura compatibilidad con versiones antiguas de sessions.json."""
    if not isinstance(store, dict):
        return {"active_id": None, "items": {}}

    items = store.get("items") or {}
    if not isinstance(items, dict):
        items = {}

    for sid, it in list(items.items()):
        if not isinstance(it, dict):
            continue

        it["created_at"] = _to_float_ts(it.get("created_at", 0))
        it["updated_at"] = _to_float_ts(it.get("updated_at", 0))

        try:
            it["top_k"] = int(it.get("top_k", DEFAULT_TOP_K))
        except Exception:
            it["top_k"] = DEFAULT_TOP_K

        it["title"] = _sanitize_title(it.get("title") or "Conversación")
        it["history"] = _normalize_history_to_messages(it.get("history") or [])
        items[sid] = it

    # Recorte opcional (rendimiento en entornos con muchas conversaciones)
    if MAX_SESSIONS and len(items) > MAX_SESSIONS:
        ordered = sorted(
            items.items(),
            key=lambda kv: _to_float_ts((kv[1] or {}).get("updated_at", 0)),
            reverse=True,
        )
        items = dict(ordered[:MAX_SESSIONS])

    store["items"] = items

    active_id = store.get("active_id")
    if active_id not in items:
        store["active_id"] = next(iter(items.keys()), None)

    return store


# -----------------------------------------------------------------------------
# Dropdown robusto: evita problemas entre versiones de Gradio (label/value)
# -----------------------------------------------------------------------------

def _choice_for_sid(store: Dict[str, Any], sid: str) -> str:
    it = (store.get("items") or {}).get(sid) or {}
    title = _sanitize_title(it.get("title") or "Conversación")
    return f"{title}  [{sid[:8]}]"


_SID_RE = re.compile(r"\[([0-9a-fA-F]{8})\]\s*$")


def _sid_from_choice(choice: str, store: Dict[str, Any]) -> Optional[str]:
    if not choice:
        return None
    m = _SID_RE.search(str(choice).strip())
    if not m:
        # Algunas versiones podrían devolver el sid directamente.
        if choice in (store.get("items") or {}):
            return choice
        return None
    pref = m.group(1).lower()
    for sid in (store.get("items") or {}).keys():
        if str(sid).lower().startswith(pref):
            return sid
    return None


# -----------------------------------------------------------------------------
# Inicialización del motor RAG (LLM + embeddings + vector store)
# -----------------------------------------------------------------------------
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    raise RuntimeError("Falta GOOGLE_API_KEY en tu entorno o en el .env")


def _make_embed_model(model_name: str):
    """Compatibilidad con distintas firmas del wrapper de embeddings."""
    try:
        return GoogleGenAIEmbedding(model_name=model_name, api_key=API_KEY)
    except TypeError:
        return GoogleGenAIEmbedding(model=model_name, api_key=API_KEY)


Settings.llm = GoogleGenAI(
    model=GEMINI_MODEL,
    api_key=API_KEY,
    temperature=0.1,
    max_output_tokens=2048,
)
Settings.embed_model = _make_embed_model(GEMINI_EMBED_MODEL)

_chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)
_chroma_collection = _chroma_client.get_or_create_collection(COLLECTION_NAME)
_vector_store = ChromaVectorStore(chroma_collection=_chroma_collection)
_index = VectorStoreIndex.from_vector_store(_vector_store)


def build_engine(similarity_top_k: int):
    return _index.as_chat_engine(
        chat_mode="context",
        similarity_top_k=int(similarity_top_k),
        system_prompt=INSTRUCCIONES,
        verbose=False,
    )


def format_sources(resp, max_sources: int = 8) -> str:
    """Formatea una lista compacta de fuentes a partir de source_nodes."""
    nodes = getattr(resp, "source_nodes", None) or []
    out = []
    seen = set()

    for n in nodes:
        md = getattr(n, "metadata", None) or {}
        titulo = (
            md.get("titulo") or md.get("title") or md.get("file_name") or "Documento"
        )
        origen = (
            md.get("origen")
            or md.get("url")
            or md.get("source")
            or md.get("file_path")
            or ""
        )
        tipo = md.get("tipo") or md.get("type") or ""

        key = (titulo, origen, tipo)
        if key in seen:
            continue
        seen.add(key)

        tag = f"[{str(tipo).upper()}] " if tipo else ""
        out.append(f"- {tag}{titulo}" + (f" (Ref: {origen})" if origen else ""))

        if len(out) >= max_sources:
            break

    if not out:
        return ""
    return "\n\n**Fuentes (top):**\n" + "\n".join(out)


# -----------------------------------------------------------------------------
# Sesiones: persistencia local en ui_sessions/sessions.json
# -----------------------------------------------------------------------------

def _ensure_sessions_dir():
    SESSIONS_DIR.mkdir(parents=True, exist_ok=True)


def _ensure_exports_dir():
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _load_sessions() -> Dict[str, Any]:
    _ensure_sessions_dir()
    if not SESSIONS_FILE.exists():
        return {"active_id": None, "items": {}}
    try:
        raw = json.loads(SESSIONS_FILE.read_text(encoding="utf-8"))
        store = _migrate_sessions_store(raw)
        _save_sessions(store)  # persistir el formato migrado
        return store
    except Exception:
        return {"active_id": None, "items": {}}


def _save_sessions(data: Dict[str, Any]) -> None:
    _ensure_sessions_dir()
    tmp = SESSIONS_FILE.with_suffix(".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(SESSIONS_FILE)


def _new_session(top_k: int) -> Tuple[str, Dict[str, Any]]:
    sid = str(uuid.uuid4())
    item = {
        "title": "Nueva conversación",
        "top_k": int(top_k),
        "history": [],
        "created_at": _now_ts(),
        "updated_at": _now_ts(),
    }
    return sid, item


def _session_choices(store: Dict[str, Any]) -> List[str]:
    """Devuelve las opciones del dropdown ordenadas por último uso."""
    items = store.get("items", {}) or {}
    ordered = sorted(
        items.items(),
        key=lambda kv: (
            _to_float_ts((kv[1] or {}).get("updated_at", 0))
            if isinstance(kv[1], dict)
            else 0.0
        ),
        reverse=True,
    )
    return [_choice_for_sid(store, sid) for sid, _ in ordered]


def _auto_title_if_needed(item: Dict[str, Any]) -> None:
    """Asigna un título automático usando el primer mensaje del usuario."""
    if (item.get("title") or "") != "Nueva conversación":
        return
    for msg in item.get("history") or []:
        role, content_raw = _extract_role_content(msg)
        if role == "user":
            text = _content_to_text(content_raw).strip()
            if text:
                item["title"] = _sanitize_title(
                    (text[:42] + "…") if len(text) > 42 else text
                )
            return


def init_app_state() -> Dict[str, Any]:
    """Estado inicial de la app (store + engine)."""
    store = _load_sessions()
    if not store.get("active_id") or store["active_id"] not in store.get("items", {}):
        sid, item = _new_session(DEFAULT_TOP_K)
        store.setdefault("items", {})[sid] = item
        store["active_id"] = sid
        _save_sessions(store)
    return {"store": store, "engine": None}


def get_active(store: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    sid = store.get("active_id")
    item = (store.get("items", {}) or {}).get(sid) or {}
    return sid, item


def _reset_engine(app_state: Dict[str, Any]):
    """Fuerza la recreación del chat engine en la siguiente interacción."""
    app_state["engine"] = None
    return app_state


# -----------------------------------------------------------------------------
# Header dinámico según haya historial o no
# -----------------------------------------------------------------------------

def header_initial() -> str:
    return """
    <div class="hdr">
      <div class="hdr-title">Hola, Saúl. ¿Todo listo para empezar?</div>
      <div class="hdr-sub">Asistente RAG</div>
    </div>
    """


def header_compact() -> str:
    return """
    <div class="hdr compact">
      <div class="hdr-title">Asistente RAG</div>
      <div class="hdr-sub">ChromaDB + LlamaIndex</div>
    </div>
    """


def _header_for_history(history: List[Dict[str, str]]) -> str:
    return header_initial() if not history else header_compact()


# -----------------------------------------------------------------------------
# Callbacks de UI (sesiones y configuración)
# -----------------------------------------------------------------------------

def on_select_session(selected_choice: str, app_state: Dict[str, Any]):
    app_state = app_state or init_app_state()
    store = _load_sessions()  # re-sincroniza estado desde disco
    app_state["store"] = store

    sid = _sid_from_choice(selected_choice, store)
    if sid and sid in store.get("items", {}):
        store["active_id"] = sid
        store["items"][sid]["updated_at"] = _now_ts()
        _save_sessions(store)

    sid, item = get_active(store)
    app_state = _reset_engine(app_state)

    history = _normalize_history_to_messages(item.get("history") or [])
    top_k = int(item.get("top_k", DEFAULT_TOP_K))

    return (
        gr.update(value=history),
        gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid)),
        top_k,
        _header_for_history(history),
        app_state,
    )


def on_new_session(top_k: int, app_state: Dict[str, Any]):
    app_state = app_state or init_app_state()
    store = _load_sessions()
    app_state["store"] = store

    sid, item = _new_session(int(top_k))
    store["items"][sid] = item
    store["active_id"] = sid
    _save_sessions(store)

    app_state = _reset_engine(app_state)

    return (
        gr.update(value=[]),
        gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid)),
        header_initial(),
        app_state,
    )


def on_delete_session(app_state: Dict[str, Any]):
    app_state = app_state or init_app_state()
    store = _load_sessions()
    app_state["store"] = store

    sid, _ = get_active(store)
    if sid and sid in store.get("items", {}):
        del store["items"][sid]

    if not store["items"]:
        nsid, nitem = _new_session(DEFAULT_TOP_K)
        store["items"][nsid] = nitem
        store["active_id"] = nsid
    else:
        ordered = sorted(
            store["items"].items(),
            key=lambda kv: _to_float_ts((kv[1] or {}).get("updated_at", 0)),
            reverse=True,
        )
        store["active_id"] = ordered[0][0]

    _save_sessions(store)
    app_state = _reset_engine(app_state)

    sid, item = get_active(store)
    history = _normalize_history_to_messages(item.get("history") or [])
    return (
        gr.update(value=history),
        gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid)),
        int(item.get("top_k", DEFAULT_TOP_K)),
        _header_for_history(history),
        app_state,
    )


def on_rename_session(new_title: str, session_choice: str, app_state: Dict[str, Any]):
    app_state = app_state or init_app_state()
    store = _load_sessions()
    app_state["store"] = store

    sid = _sid_from_choice(session_choice, store) or store.get("active_id")
    item = (store.get("items") or {}).get(sid) or {}

    t = _sanitize_title(new_title or "")
    if sid and t:
        item["title"] = t
        item["updated_at"] = _now_ts()
        store["items"][sid] = item
        store["active_id"] = sid
        _save_sessions(store)

    return (
        gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid)),
        "",
    )


def on_topk_change(top_k: int, session_choice: str, app_state: Dict[str, Any]):
    """Al cambiar top_k se limpia el chat activo para evitar inconsistencias."""
    app_state = app_state or init_app_state()
    store = _load_sessions()
    app_state["store"] = store

    sid = _sid_from_choice(session_choice, store) or store.get("active_id")
    item = (store.get("items") or {}).get(sid) or {}

    if sid:
        item["top_k"] = int(top_k)
        item["history"] = []
        item["updated_at"] = _now_ts()
        store["items"][sid] = item
        store["active_id"] = sid
        _save_sessions(store)

    app_state = _reset_engine(app_state)
    return gr.update(value=[]), header_initial(), app_state


def on_clear_chat(session_choice: str, app_state: Dict[str, Any]):
    app_state = app_state or init_app_state()
    store = _load_sessions()
    app_state["store"] = store

    sid = _sid_from_choice(session_choice, store) or store.get("active_id")
    item = (store.get("items") or {}).get(sid) or {}

    item["history"] = []
    item["updated_at"] = _now_ts()
    store["items"][sid] = item
    store["active_id"] = sid
    _save_sessions(store)

    app_state = _reset_engine(app_state)
    return gr.update(value=[]), header_initial(), app_state


def on_app_load(app_state: Dict[str, Any]):
    """Sincroniza UI desde disco al cargar/recargar (F5/Ctrl+R)."""
    app_state = app_state or init_app_state()
    store = _load_sessions()

    # Asegura active_id válido
    if not store.get("active_id") or store["active_id"] not in (store.get("items") or {}):
        sid, item = _new_session(DEFAULT_TOP_K)
        store.setdefault("items", {})[sid] = item
        store["active_id"] = sid
        _save_sessions(store)

    app_state["store"] = store
    app_state = _reset_engine(app_state)

    sid, item = get_active(store)
    history = _normalize_history_to_messages(item.get("history") or [])
    top_k_val = int(item.get("top_k", DEFAULT_TOP_K))

    return (
        gr.update(value=history),
        gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid)),
        top_k_val,
        _header_for_history(history),
        app_state,
    )


# -----------------------------------------------------------------------------
# Exportación DOCX (Markdown -> Word, conversión básica)
# -----------------------------------------------------------------------------

def _safe_filename(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"[^\w\s\-\(\)\.]", "", s, flags=re.UNICODE)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:80] if s else "conversacion"


def _add_inline_md(paragraph, text: str):
    """Parser mínimo: **negrita**, *cursiva*, `código`, [texto](url)."""
    i = 0
    n = len(text)
    while i < n:
        # Enlace
        if text[i] == "[":
            j = text.find("](", i + 1)
            k = text.find(")", j + 2) if j != -1 else -1
            if j != -1 and k != -1:
                label = text[i + 1 : j]
                url = text[j + 2 : k]
                r = paragraph.add_run(label)
                r.underline = True
                paragraph.add_run(f" ({url})")
                i = k + 1
                continue

        # Negrita
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                run = paragraph.add_run(text[i + 2 : j])
                run.bold = True
                i = j + 2
                continue

        # Código
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j != -1:
                run = paragraph.add_run(text[i + 1 : j])
                try:
                    run.font.name = "Consolas"
                except Exception:
                    pass
                i = j + 1
                continue

        # Cursiva
        if text[i] == "*":
            j = text.find("*", i + 1)
            if j != -1:
                run = paragraph.add_run(text[i + 1 : j])
                run.italic = True
                i = j + 1
                continue

        # Tramo sin tokens
        next_pos = n
        for token in ("[", "**", "`", "*"):
            p = text.find(token, i + 1)
            if p != -1:
                next_pos = min(next_pos, p)
        paragraph.add_run(text[i:next_pos])
        i = next_pos


def _md_to_docx(doc, text: str):
    lines = (text or "").splitlines()
    in_code = False
    code_buf: List[str] = []

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                try:
                    run.font.name = "Consolas"
                except Exception:
                    pass
            continue

        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(6, len(m.group(1)))
            doc.add_heading(m.group(2).strip(), level=level)
            continue

        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_md(p, m.group(1))
            continue

        m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_md(p, m.group(1))
            continue

        p = doc.add_paragraph()
        _add_inline_md(p, line)


def export_current_to_docx(session_choice: str, app_state: Dict[str, Any]):
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("Falta python-docx. Instalar con: pip install python-docx") from e

    app_state = app_state or init_app_state()
    store = _load_sessions()
    app_state["store"] = store

    sid = _sid_from_choice(session_choice, store) or store.get("active_id")
    item = (store.get("items") or {}).get(sid) or {}
    history = _normalize_history_to_messages(item.get("history") or [])

    title = _sanitize_title(item.get("title") or "Conversación")
    fname = _safe_filename(title) + f"_{sid[:8]}.docx"

    _ensure_exports_dir()
    out_path = EXPORTS_DIR / fname

    doc = Document()
    doc.add_heading(title, level=1)

    for msg in history:
        role, content_raw = _extract_role_content(msg)
        content = _content_to_text(content_raw)

        if role == "user":
            doc.add_heading("Usuario", level=2)
            _md_to_docx(doc, content)
        elif role == "assistant":
            doc.add_heading("Asistente", level=2)
            _md_to_docx(doc, content)
        else:
            continue

        doc.add_paragraph("")

    doc.save(str(out_path))
    return str(out_path)


# -----------------------------------------------------------------------------
# Chat: generator con streaming si existe; si no, animación de "..."
# -----------------------------------------------------------------------------
_EXECUTOR = ThreadPoolExecutor(max_workers=2)


def reply(
    user_msg: str,
    history,
    top_k: int,
    streaming: bool,
    session_choice: str,
    app_state: Dict[str, Any],
):
    user_msg = (user_msg or "").strip()
    history = _normalize_history_to_messages(history or [])

    if not user_msg:
        yield gr.update(value=history), gr.update(), _header_for_history(history), app_state, ""
        return

    app_state = app_state or init_app_state()
    store = _load_sessions()  # el estado se toma siempre desde disco
    app_state["store"] = store

    sid = _sid_from_choice(session_choice, store) or store.get("active_id")
    if sid and sid in store.get("items", {}):
        store["active_id"] = sid

    sid, item = get_active(store)
    top_k = int(top_k)

    # Consistencia de top_k: si cambia, se reinicia el historial
    if int(item.get("top_k", DEFAULT_TOP_K)) != top_k:
        item["top_k"] = top_k
        item["history"] = []
        history = []
        app_state = _reset_engine(app_state)

    # Crear engine si aún no existe en memoria
    if app_state.get("engine") is None:
        app_state["engine"] = build_engine(top_k)

    # Añadir mensaje del usuario + placeholder del asistente
    history.append({"role": "user", "content": user_msg})
    history.append({"role": "assistant", "content": "..."})

    # Persistencia inmediata (incluye placeholder)
    item["history"] = history
    item["updated_at"] = _now_ts()
    _auto_title_if_needed(item)
    store["items"][sid] = item
    store["active_id"] = sid
    _save_sessions(store)

    # Refrescar dropdown al inicio (por si el auto-título cambia)
    dd_update = gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid))
    yield gr.update(value=history), dd_update, header_compact(), app_state, ""

    try:
        engine = app_state["engine"]

        # Streaming real si el engine lo expone
        if streaming and hasattr(engine, "stream_chat"):
            resp = engine.stream_chat(user_msg)
            partial = ""

            gen = getattr(resp, "response_gen", None)
            if gen is None:
                final = _content_to_text(getattr(resp, "response", str(resp))).strip()
                sources = format_sources(resp)
                history[-1]["content"] = final + (("\n" + sources) if sources else "")
                yield gr.update(value=history), gr.update(), header_compact(), app_state, ""
            else:
                for chunk in gen:
                    partial += str(chunk)
                    history[-1]["content"] = partial
                    yield gr.update(value=history), gr.update(), header_compact(), app_state, ""

                sources = format_sources(resp)
                if sources:
                    history[-1]["content"] = (history[-1]["content"] or "").rstrip() + "\n" + sources
                yield gr.update(value=history), gr.update(), header_compact(), app_state, ""

        else:
            # Sin streaming: se muestra una animación breve mientras se resuelve
            dots = [".", "..", "..."]
            i = 0

            def _do_chat():
                return engine.chat(user_msg)

            fut = _EXECUTOR.submit(_do_chat)

            while not fut.done():
                history[-1]["content"] = dots[i % len(dots)]
                yield gr.update(value=history), gr.update(), header_compact(), app_state, ""
                time.sleep(0.35)
                i += 1

            resp = fut.result()
            answer = _content_to_text(getattr(resp, "response", str(resp))).strip()
            sources = format_sources(resp)
            if sources:
                answer = f"{answer}\n{sources}"
            history[-1]["content"] = answer
            yield gr.update(value=history), gr.update(), header_compact(), app_state, ""

    except Exception as e:
        history[-1]["content"] = f"❌ {type(e).__name__}: {e}\n\n```text\n{traceback.format_exc()}\n```"
        yield gr.update(value=history), gr.update(), header_compact(), app_state, ""

    # Persistencia final tras completar la respuesta
    item["history"] = history
    item["updated_at"] = _now_ts()
    _auto_title_if_needed(item)
    store["items"][sid] = item
    store["active_id"] = sid
    _save_sessions(store)

    # Refrescar dropdown al final (por si cambió el título)
    dd_update_end = gr.update(choices=_session_choices(store), value=_choice_for_sid(store, sid))
    yield gr.update(value=history), dd_update_end, header_compact(), app_state, ""


# -----------------------------------------------------------------------------
# UI: estilos y comportamiento de teclado
# -----------------------------------------------------------------------------
CSS = """
:root { --radius: 14px; }

.gradio-container { max-width: 1500px !important; margin: 0 auto !important; }
footer { display: none !important; }

/* Ocultar etiquetas tipo "Chatbot", "Textbox", etc. */
.label, .wrap .label, .block .label { display: none !important; }

/* Header */
.hdr { text-align:center; margin: 10px 0 10px 0; }
.hdr.compact { margin-bottom: 8px; }
.hdr-title { font-size: 34px; font-weight: 750; line-height: 1.15; }
.hdr-sub { opacity: .75; margin-top: 6px; font-size: 14px; }

/* Layout principal */
#layout { align-items: stretch; gap: 14px; flex-wrap: nowrap !important; }

/* Sidebar */
#sidebar {
  height: calc(100vh - 140px);
  position: sticky;
  top: 10px;
  border-radius: var(--radius);
  min-width: 300px !important;
}

/* Columna del chat */
#chatcol { min-width: 980px !important; }
#chatbox {
  height: calc(100vh - 260px);
  min-height: 600px;
  border-radius: var(--radius);
}

/* Composer */
#composer { gap: 10px; align-items: stretch; flex-wrap: nowrap !important; }
#msg { border-radius: var(--radius); min-height: 96px; }

/* Ocultar botón compartir del chat (distintos selectores por compatibilidad) */
#chatbox button[aria-label="Share"],
#chatbox button[aria-label="Compartir"],
#chatbox button[title="Share"],
#chatbox button[title="Compartir"],
#chatbox .share-button,
#chatbox .gr-share-btn {
  display: none !important;
}
"""


# JS: Enter envía; Shift+Enter inserta nueva línea.
# Se reintenta varias veces porque Gradio puede re-renderizar componentes.
JS_ON_LOAD = r"""
() => {
  const attach = () => {
    const ta = document.querySelector('#msg textarea');
    if (!ta) return false;
    if (ta.dataset.enterSendAttached === "1") return true;

    const clickSend = () => {
      const btn = document.querySelector('#send_btn button') || document.querySelector('#send_btn');
      if (btn) btn.click();
    };

    ta.addEventListener('keydown', (e) => {
      if (e.isComposing) return;                // IME
      if (e.key !== 'Enter') return;
      if (e.shiftKey) return;                   // Shift+Enter -> newline normal
      e.preventDefault();                       // Enter -> enviar
      clickSend();
    });

    ta.dataset.enterSendAttached = "1";
    return true;
  };

  // Intenta varias veces (Gradio re-renderiza)
  let tries = 0;
  const it = setInterval(() => {
    tries += 1;
    const ok = attach();
    if (ok || tries > 30) clearInterval(it);
  }, 250);
}
"""


def _make_chatbot(**kwargs):
    """Crea un Chatbot intentando desactivar el botón de compartir si existe la opción."""
    try:
        return gr.Chatbot(show_share_button=False, **kwargs)
    except TypeError:
        return gr.Chatbot(**kwargs)


with gr.Blocks(title="Asistente RAG") as demo:
    initial_state = init_app_state()
    initial_store = initial_state["store"]
    initial_sid = initial_store["active_id"]
    initial_item = initial_store["items"][initial_sid]

    initial_history = _normalize_history_to_messages(initial_item.get("history") or [])

    app_state = gr.State(initial_state)

    header = gr.HTML(_header_for_history(initial_history))

    with gr.Row(elem_id="layout"):
        # Sidebar
        with gr.Column(scale=1, min_width=300, elem_id="sidebar"):
            gr.Markdown("### Conversaciones")

            session_dd = gr.Dropdown(
                label="",
                show_label=False,
                choices=_session_choices(initial_store),
                value=_choice_for_sid(initial_store, initial_sid),
                interactive=True,
            )

            with gr.Row():
                btn_new = gr.Button("Nueva")
                btn_del = gr.Button("Borrar")

            title_box = gr.Textbox(
                label="",
                show_label=False,
                placeholder="Renombrar… (Enter)",
            )

            gr.Markdown("### Ajustes")
            top_k = gr.Dropdown(
                label="similarity_top_k",
                show_label=True,
                choices=TOP_K_OPTIONS,
                value=int(initial_item.get("top_k", DEFAULT_TOP_K)),
                info="5 = más preciso/corto · 10 = balance · 20 = recupera más",
                interactive=True,
            )
            streaming = gr.Checkbox(
                label="Streaming (si está disponible)",
                value=True,
            )

            btn_clear = gr.Button("Limpiar chat")
            btn_export = gr.Button("Exportar .docx")
            export_file = gr.File(label="", show_label=False)

        # Chat principal
        with gr.Column(scale=5, elem_id="chatcol"):
            chatbot = _make_chatbot(
                label="",
                show_label=False,
                elem_id="chatbox",
                value=initial_history,
            )

            with gr.Row(elem_id="composer"):
                msg = gr.Textbox(
                    placeholder="Enter para enviar · Shift+Enter nueva línea",
                    label="",
                    show_label=False,
                    elem_id="msg",
                    lines=3,
                    scale=12,
                )
                send = gr.Button("Enviar", variant="primary", scale=2, elem_id="send_btn")

    # Al cargar/recargar: refresca sesiones desde disco y activa el binding de teclado.
    demo.load(
        fn=on_app_load,
        inputs=[app_state],
        outputs=[chatbot, session_dd, top_k, header, app_state],
        js=JS_ON_LOAD,
    )

    # Wiring
    session_dd.change(
        fn=on_select_session,
        inputs=[session_dd, app_state],
        outputs=[chatbot, session_dd, top_k, header, app_state],
    )

    btn_new.click(
        fn=on_new_session,
        inputs=[top_k, app_state],
        outputs=[chatbot, session_dd, header, app_state],
    )

    btn_del.click(
        fn=on_delete_session,
        inputs=[app_state],
        outputs=[chatbot, session_dd, top_k, header, app_state],
    )

    title_box.submit(
        fn=on_rename_session,
        inputs=[title_box, session_dd, app_state],
        outputs=[session_dd, title_box],
    )

    top_k.change(
        fn=on_topk_change,
        inputs=[top_k, session_dd, app_state],
        outputs=[chatbot, header, app_state],
    )

    btn_clear.click(
        fn=on_clear_chat,
        inputs=[session_dd, app_state],
        outputs=[chatbot, header, app_state],
    )

    btn_export.click(
        fn=export_current_to_docx,
        inputs=[session_dd, app_state],
        outputs=[export_file],
    )

    send.click(
        fn=reply,
        inputs=[msg, chatbot, top_k, streaming, session_dd, app_state],
        outputs=[chatbot, session_dd, header, app_state, msg],
    )

    # Mantener submit por compatibilidad: el envío principal se gestiona vía JS.
    msg.submit(
        fn=reply,
        inputs=[msg, chatbot, top_k, streaming, session_dd, app_state],
        outputs=[chatbot, session_dd, header, app_state, msg],
    )

if __name__ == "__main__":
    demo.launch(
        server_name="127.0.0.1",
        server_port=7860,
        inbrowser=True,
        show_error=True,
        theme=gr.themes.Soft(),
        css=CSS,
    )
